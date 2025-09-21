import os
from typing import List
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from ..config import settings


class DocumentProcessor:
    def __init__(self, chunk_size: int = None, chunk_overlap: int = None):
        # Use configuration values if not provided
        chunk_size = chunk_size or settings.CHUNK_SIZE
        chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def process_document(self, file_path: str) -> List[LangchainDocument]:
        """Process a document and return text chunks"""
        try:
            # Extract text based on file type
            if file_path.lower().endswith('.pdf'):
                text = self._extract_pdf_text(file_path)
            elif file_path.lower().endswith(('.docx', '.doc')):
                text = self._extract_word_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_path}")
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Convert to Langchain documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc = LangchainDocument(
                    page_content=chunk,
                    metadata={
                        "source": os.path.basename(file_path),
                        "chunk_id": i,
                        "file_type": file_path.split('.')[-1].lower()
                    }
                )
                documents.append(doc)
            
            return documents
        
        except Exception as e:
            raise Exception(f"Error processing document {file_path}: {str(e)}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def _extract_word_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
